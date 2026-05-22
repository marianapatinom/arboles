import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Establece el color de fondo de una celda de tabla en Word."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Añade padding interno (márgenes) a una celda para una mejor legibilidad."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, color="CCCCCC", sz="4", val="single"):
    """Establece bordes limpios y discretos en una celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def create_premium_docx(md_path, docx_path, img_dir):
    doc = Document()
    
    # Configuración de márgenes estándar de la página (2.54 cm / 1 pulgada)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Colores corporativos (Teal oscuro y Teal claro)
    PRIMARY_COLOR = RGBColor(11, 59, 58)    # #0B3B3A
    SECONDARY_COLOR = RGBColor(43, 122, 120) # #2B7A78
    BODY_COLOR = RGBColor(51, 51, 51)        # #333333
    
    # Estilo Normal (Cuerpo de texto)
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = BODY_COLOR
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 1. PORTADA ACADÉMICA (Diseño Corporativo)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(80)
    title_p.paragraph_format.space_after = Pt(12)
    title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("MEMORIA TÉCNICA\n")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY_COLOR
    
    run_main_title = title_p.add_run("CLASIFICACIÓN TAXONÓMICA DE ESPECIES DE FLORES DEL GÉNERO IRIS MEDIANTE EL ALGORITMO DE ÁRBOL DE DECISIÓN ID3")
    run_main_title.font.name = 'Calibri'
    run_main_title.font.size = Pt(22)
    run_main_title.font.bold = True
    run_main_title.font.color.rgb = PRIMARY_COLOR
    
    # Línea decorativa
    border_p = doc.add_paragraph()
    border_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border_p.paragraph_format.space_after = Pt(150)
    run_line = border_p.add_run("____________________________________________________")
    run_line.font.color.rgb = SECONDARY_COLOR
    run_line.font.bold = True
    
    # Datos de autores y fecha
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.line_spacing = 1.3
    
    add_meta(meta_p, "Autor: ", SECONDARY_COLOR, True)
    add_meta(meta_p, "Agente de Construcción de Proyectos de IA\n", BODY_COLOR, False)
    
    add_meta(meta_p, "Institución: ", SECONDARY_COLOR, True)
    add_meta(meta_p, "Proyecto IA - Talento Tech\n", BODY_COLOR, False)
    
    add_meta(meta_p, "Metodología: ", SECONDARY_COLOR, True)
    add_meta(meta_p, "CRISP-ML (Cross-Industry Standard Process)\n", BODY_COLOR, False)
    
    add_meta(meta_p, "Fecha: ", SECONDARY_COLOR, True)
    add_meta(meta_p, "22 de Mayo de 2026", BODY_COLOR, False)
    
    # Salto de página para iniciar el contenido
    doc.add_page_break()
    
    # Leer archivo Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_rows = []
    
    # Control de imágenes a insertar después de secciones específicas
    pending_images = {
        "4. ANÁLISIS DEL CONJUNTO DE DATOS": "GraficoBarras.png",
        "5. DESARROLLO DEL MODELO DE MODELADO": "VisualizaacionArbol.png",
        "6. EVALUACIÓN CIENTÍFICA DE RESULTADOS": "MatrizConfusion.png"
    }
    
    current_section = ""
    
    for i, line in enumerate(lines):
        line_str = line.strip()
        
        # Saltar metadatos iniciales ya renderizados en la portada
        if i < 15 and (line_str.startswith("---") or line_str.startswith("**Autores:**") or line_str.startswith("**Institución:**") or line_str.startswith("**Fecha:**") or "MEMORIA TÉCNICA" in line_str):
            continue
            
        # Detectar fin de tabla y renderizarla si existe
        if not line_str.startswith("|") and in_table:
            render_docx_table(doc, table_rows, PRIMARY_COLOR, SECONDARY_COLOR)
            in_table = False
            table_rows = []
            
        # 2. PROCESAR ENCABEZADOS Y ELEMENTOS DE MARKDOWN
        if line_str.startswith("# "):
            header_text = line_str.lstrip("# ").strip()
            # Si era el título principal, lo ignoramos porque ya está en la portada
            if "MEMORIA TÉCNICA" in header_text.upper():
                continue
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(header_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = PRIMARY_COLOR
            current_section = header_text
            
        elif line_str.startswith("## "):
            header_text = line_str.lstrip("## ").strip()
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(header_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = SECONDARY_COLOR
            
        elif line_str.startswith("### "):
            header_text = line_str.lstrip("### ").strip()
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(2)
            h.paragraph_format.keep_with_next = True
            run = h.add_run(header_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = SECONDARY_COLOR
            
        elif line_str.startswith("|"):
            in_table = True
            table_rows.append(line_str)
            
        elif line_str.startswith("* ") or line_str.startswith("- "):
            bullet_text = line_str.lstrip("*- ").strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            parse_formatted_text(p, bullet_text, SECONDARY_COLOR)
            
        elif line_str.startswith("1. ") or line_str.startswith("2. ") or line_str.startswith("3. ") or line_str.startswith("4. ") or line_str.startswith("5. ") or line_str.startswith("6. ") or line_str.startswith("7. "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            parse_formatted_text(p, line_str, SECONDARY_COLOR)
            
        elif line_str.startswith("---") or line_str == "":
            continue
            
        elif line_str.startswith("$$") or line_str.startswith("$$\\text"):
            # Ecuaciones matemáticas
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line_str.replace("$$", "").strip())
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = PRIMARY_COLOR
            
        elif line_str.startswith("```"):
            # Bloque de código o diagrama (los ignoramos si son de tipo mermaid)
            continue
            
        else:
            p = doc.add_paragraph()
            parse_formatted_text(p, line_str, SECONDARY_COLOR)
            
        # 3. VERIFICAR INSERCIÓN DE IMÁGENES
        # Al finalizar el análisis de una sección, insertamos la imagen correspondiente
        # Vemos si la siguiente línea empieza con "# " (cambio de sección) o si es el final
        next_is_header = False
        if i + 1 < len(lines):
            next_line = lines[i+1].strip()
            if next_line.startswith("# "):
                next_is_header = True
                
        if (next_is_header or i == len(lines) - 1) and current_section in pending_images:
            img_name = pending_images[current_section]
            img_path = os.path.join(img_dir, img_name)
            if os.path.exists(img_path):
                # Insertar salto para espaciado
                p_spacer = doc.add_paragraph()
                p_spacer.paragraph_format.space_before = Pt(6)
                
                # Insertar imagen centrada
                p_img = doc.add_paragraph()
                p_img.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_after = Pt(4)
                p_img.add_run().add_picture(img_path, width=Inches(5.0))
                
                # Pie de foto
                p_cap = doc.add_paragraph()
                p_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(14)
                run_cap = p_cap.add_run(f"Figura: Diagrama de soporte técnico - {img_name.replace('.png', '')}")
                run_cap.font.name = 'Calibri'
                run_cap.font.size = Pt(9.5)
                run_cap.font.italic = True
                run_cap.font.color.rgb = SECONDARY_COLOR
                
            del pending_images[current_section]

    # Guardar documento
    try:
        doc.save(docx_path)
        print(f"Documento premium generado exitosamente en: {docx_path}")
    except PermissionError:
        print(f"\n[ERROR DE PERMISO] No se pudo guardar el archivo en:\n{docx_path}")
        print("Esto ocurre porque el documento está abierto en Microsoft Word o bloqueado por otro proceso.")
        print("Por favor, CIERRE Microsoft Word y ejecute nuevamente el script.\n")
        sys.exit(1)

def add_meta(paragraph, label, color, is_bold):
    run = paragraph.add_run(label)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = is_bold
    run.font.color.rgb = color

def parse_formatted_text(paragraph, text, accent_rgb):
    """Parsea texto markdown simple (negrita ** y código `) a estilos de Word."""
    parts = text.split("**")
    for idx, part in enumerate(parts):
        is_bold = (idx % 2 == 1)
        subparts = part.split("`")
        for s_idx, subpart in enumerate(subparts):
            is_code = (s_idx % 2 == 1)
            run = paragraph.add_run(subpart)
            run.font.name = 'Consolas' if is_code else 'Calibri'
            run.font.size = Pt(10) if is_code else Pt(11)
            run.font.bold = is_bold
            if is_code:
                run.font.color.rgb = accent_rgb

def render_docx_table(doc, rows, primary_color, secondary_color):
    """Parsea un bloque de filas markdown en una tabla estilizada de python-docx."""
    # Eliminar fila divisoria si existe (ej: | :--- | :---: |)
    filtered_rows = []
    for r in rows:
        if "---" not in r:
            filtered_rows.append(r)
            
    if not filtered_rows:
        return
        
    # Obtener dimensiones
    header_cols = [c.strip() for c in filtered_rows[0].split("|")[1:-1]]
    num_cols = len(header_cols)
    num_rows = len(filtered_rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Procesar fila de cabecera
    hdr_cells = table.rows[0].cells
    for c_idx, text in enumerate(header_cols):
        cell = hdr_cells[c_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, "0B3B3A") # Navy Blue
        set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
        set_cell_borders(cell, color="444444", sz="6")
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255) # Blanco
        
    # Procesar filas de datos
    for r_idx in range(1, num_rows):
        cols = [c.strip() for c in filtered_rows[r_idx].split("|")[1:-1]]
        row_cells = table.rows[r_idx].cells
        
        # Color cebra alternado para filas
        bg_color = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        
        for c_idx in range(min(num_cols, len(cols))):
            cell = row_cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            set_cell_borders(cell, color="DDDDDD", sz="4")
            p = cell.paragraphs[0]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(0)
            parse_formatted_text(p, cols[c_idx], secondary_color)
            
    # Añadir un espacio después de la tabla
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(8)

if __name__ == '__main__':
    workspace_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    default_md = os.path.join(workspace_dir, "Proyecto", "Memoria_Tecnica.md")
    default_docx = os.path.join(workspace_dir, "Proyecto", "Material_1_Proyecto_IA_Talento_Tech.docx")
    default_img_dir = os.path.join(workspace_dir, "Imágenes")
    
    md_file = sys.argv[1] if len(sys.argv) > 1 else default_md
    docx_file = sys.argv[2] if len(sys.argv) > 2 else default_docx
    img_dir = sys.argv[3] if len(sys.argv) > 3 else default_img_dir
    
    create_premium_docx(md_file, docx_file, img_dir)
